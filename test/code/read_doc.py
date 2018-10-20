#读取docx中的文本代码示例
import docx
import re
#获取文档对象
file=docx.Document("../data/常用电力词汇.docx")
#print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

#输出每一段的内容
ww = []
for para in file.paragraphs:
    ww.append(para.text)
    # print(para.text)
print(len(ww))
print(type(ww[1]))

ww= [''.join(re.findall(u'[\u4e00-\u9fff]+',i))for i in ww]
print(ww)

# #输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)
